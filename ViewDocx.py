from docx import Document
from docx.enum.section import WD_SECTION
from docx.shared import Pt
import sys
import os

def watermarkImport(dirPath_input, inputPath_input):

    section = dirPath_input.sections[0]
    footer = section.footer

    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = 2

    # Add text as watermark in the footer
    run = paragraph.add_run()
    font = run.font
    font.size = Pt(10)
    font.bold = True
    run.text = inputPath_input

    return doc




if __name__ == "__main__":


    if len(sys.argv) != 3:
        print("Usage: python script.py doc_path inputPath")
        sys.exit(1)


    dirPath = sys.argv[1]
    inputPath = sys.argv[2]

    doc = Document(dirPath)


    doc = watermarkImport(doc, inputPath)

    doc.save(dirPath)





