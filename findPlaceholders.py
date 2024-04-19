import json
import re
from docx import Document
import sys

def find_placeholders(doc_path):
    """
    Find all unique placeholders in the format ${key} in the document (including in text, tables,
    headers, footers, and nested tables) and return them as a JSON string. Each key maps to an empty list,
    to match the expected input format of the jsonToVector C++ function.
    """
    doc = Document(doc_path)
    placeholder_pattern = re.compile(r"\$\{(.*?)\}")
    placeholders = {}

    # Function to add matches from a given text
    def add_matches(text):
        matches = placeholder_pattern.findall(text)
        for match in matches:
            if match not in placeholders:
                placeholders[match] = []

    # Function to search recursively in table cells, including nested tables
    def search_cells(cell):
        add_matches(cell.text)
        # Check for nested tables within the cell
        for table in cell.tables:
            for row in table.rows:
                for cell in row.cells:
                    search_cells(cell)

    # Search in paragraphs
    for paragraph in doc.paragraphs:
        add_matches(paragraph.text)

    # Search in table cells, including nested tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                search_cells(cell)

    # Search in headers and footers
    for section in doc.sections:
        for header in section.header.paragraphs:
            add_matches(header.text)
        for footer in section.footer.paragraphs:
            add_matches(footer.text)

    return json.dumps(placeholders, indent=4)

if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: python find_placeholders.py doc_path")
        sys.exit(1)

    doc_path = sys.argv[1]
    placeholders_json = find_placeholders(doc_path)
    print(placeholders_json)
