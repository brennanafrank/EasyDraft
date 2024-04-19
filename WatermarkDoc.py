from docx import Document
from docx.shared import Inches
import os
import sys

def watermarkImport(input_docx, input_image):

    doc = Document(input_docx)

    footer = doc.sections[0].footer

    if not footer.paragraphs:
        footer.add_paragraph()

    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = 2

    run = footer_paragraph.add_run()

    run.add_picture(input_image, width=Inches(0.5), height=Inches(0.5))

    doc.save(input_docx)

if __name__ == "__main__":


    if len(sys.argv) != 3:
        print("Usage: python script.py doc_path image")
        sys.exit(1)


    dirPath = sys.argv[1]
    imagePath = sys.argv[2]

    watermarkImport(dirPath, imagePath)






